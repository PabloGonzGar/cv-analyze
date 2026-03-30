import docx
import io
import logging

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """
    Extrae texto de un archivo DOCX dada su ruta.
    """
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logger.error(f"Error al extraer texto del DOCX {file_path}: {e}")
        return ""

def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """
    Extrae texto de un archivo DOCX en formato de bytes (útil para archivos subidos vía API).
    """
    try:
        # Se requiere envolver los bytes en un objeto BytesIO para que docx lo trate como archivo
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logger.error(f"Error al extraer texto del DOCX desde bytes: {e}")
        return ""
